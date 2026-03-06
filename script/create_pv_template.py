from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy

# Open the template - this preserves header/footer/logo/styles
doc = Document("/Users/epinault/tmp/CSE/template CSE.docx")

# Clear all body content (paragraphs + tables) but keep section/header/footer
body = doc.element.body
for child in list(body):
    if child.tag != qn('w:sectPr'):
        body.remove(child)

# ===== HELPER FUNCTIONS =====
def add_empty_para(doc, count=1):
    for _ in range(count):
        doc.add_paragraph('', style='Normal')

def add_orange_title_line(doc, text, size=Pt(36)):
    """Add a run in orange (FF8500) to a paragraph"""
    p = doc.add_paragraph('', style='Normal')
    run = p.add_run(text)
    run.font.size = size
    run.font.color.rgb = RGBColor(0xFF, 0x85, 0x00)
    return p

def add_heading1(doc, text):
    h = doc.add_heading(text, level=1)
    return h

def add_heading2(doc, text):
    h = doc.add_heading(text, level=2)
    return h

def add_no_spacing(doc, text, bold=False, underline=False):
    p = doc.add_paragraph('', style='No Spacing')
    run = p.add_run(text)
    run.bold = bold
    run.underline = underline
    return p

def add_normal(doc, text=''):
    p = doc.add_paragraph(text, style='Normal')
    return p

def add_normal_mixed(doc, parts):
    """parts: list of (text, bold, italic, underline) tuples"""
    p = doc.add_paragraph('', style='Normal')
    for part in parts:
        text = part[0]
        bold = part[1] if len(part) > 1 else False
        italic = part[2] if len(part) > 2 else False
        underline = part[3] if len(part) > 3 else False
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph('', style='Normal')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    run = p.add_run('•  ' + text)
    return p

# ===== BUILD THE DOCUMENT =====

# --- TITLE PAGE ---
add_empty_para(doc, 9)

# Title in orange
add_orange_title_line(doc, '2026-02-10')
add_orange_title_line(doc, 'PV')
add_orange_title_line(doc, 'Réunion CSE')
add_orange_title_line(doc, 'Direction')

add_empty_para(doc, 2)

# --- Document properties section ---
p = add_normal(doc)
run = p.add_run('Propriétés du document')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

add_empty_para(doc)

# Properties table
prop_table = doc.add_table(rows=2, cols=2)
prop_table.style = 'List Table 4 Accent 1'
prop_table.rows[0].cells[0].text = 'Document Name'
c = prop_table.rows[0].cells[1]
c.text = ''
cp = c.paragraphs[0]
for text in ['2026-02-10', ' ', 'PV', ' ', 'Réunion CSE', '-', 'Direction']:
    run = cp.add_run(text)
    run.font.size = Pt(14)
prop_table.rows[1].cells[0].text = 'Document Owner'
prop_table.rows[1].cells[1].text = 'CSE Netatmo'

add_empty_para(doc)

# Change tracker
p = add_normal(doc)
run = p.add_run('Change tracker')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

ct_table = doc.add_table(rows=2, cols=5)
ct_table.style = 'List Table 4 Accent 1'
headers = ['Version', 'Date', 'Person', 'Status / Comment', 'Reviewed by / Co-writers']
for i, h in enumerate(headers):
    ct_table.rows[0].cells[i].text = h
ct_table.rows[1].cells[0].text = '1.0'
ct_table.rows[1].cells[1].text = '10/02/2026'
ct_table.rows[1].cells[2].text = ''
ct_table.rows[1].cells[3].text = '2026-02-10-PV Réunion CSE Direction'
ct_table.rows[1].cells[4].text = 'Team CSE + Direction'
# Set small font for data row
for ci in range(5):
    for p in ct_table.rows[1].cells[ci].paragraphs:
        for r in p.runs:
            r.font.size = Pt(10)

add_empty_para(doc)

# Page break
doc.add_page_break()

# ===== PRESENTS =====
add_heading1(doc, 'Présent·es :')

add_no_spacing(doc, '')
add_no_spacing(doc, 'Présents :', bold=True, underline=True)

presents_direction = [
    'Alexandre MENU, Président',
    'Sarah ETCHEVERRY, Direction',
    'Manuela BOURBOULON, Direction',
]
presents_cse = [
    'Erwan PINAULT, Secrétaire',
    'Alice LEBOIS',
    'Emile GAULTIER',
    'Céline BENSAHLA-TANI',
    'Edouard CHARPENTIER',
    'Michael LEYSSENNE',
    'Giada-Maria PIAZZA',
    'Nadège (membre CSE)',
    'Antoine BESSIN (participation partielle)',
]
for name in presents_direction + presents_cse:
    add_no_spacing(doc, name)

add_no_spacing(doc, '')
add_no_spacing(doc, 'Invités :', bold=True, underline=True)
invites = [
    'Loïc (Responsable RH R&D France) - en visioconférence',
    'Représentant LHH (cabinet d\'accompagnement / reclassement)',
]
for name in invites:
    add_no_spacing(doc, name)

add_no_spacing(doc, '')
add_no_spacing(doc, 'Absents :', bold=True, underline=True)
absents = [
    'Christophe BOLUSSET (en arrêt)',
]
for name in absents:
    add_no_spacing(doc, name)

add_empty_para(doc)

# ===== ORDRE DU JOUR =====
add_heading1(doc, 'Ordre du jour')

items_odj = [
    'Approbation du PV de la réunion précédente',
    'Consultation sur le plan d\'actionnariat salarié 2026',
    'Consultation sur la rupture conventionnelle d\'une salariée protégée',
    'Suivi de l\'accompagnement des congés de mobilité (LHH)',
    'Point sur le projet Career Pass (carrières et parcours professionnels)',
    'Situation de l\'équipe Office Management – Alerte RPS',
    'Questions diverses',
]
for i, item in enumerate(items_odj, 1):
    add_normal(doc, f'{i}. {item}')

add_empty_para(doc)

# ===== POINT 1 =====
add_heading1(doc, 'Point 1 – Approbation du PV de la réunion précédente')

add_normal(doc, 'Le procès-verbal de la réunion précédente est soumis à l\'approbation des membres du CSE.')
add_empty_para(doc)
add_normal_mixed(doc, [
    ('Décision : ', True),
    ('Le PV est approuvé.', False),
])

add_empty_para(doc)

# ===== POINT 2 =====
add_heading1(doc, 'Point 2 – Consultation sur le plan d\'actionnariat salarié 2026')

add_heading2(doc, 'Contexte')
add_normal(doc,
    'La Direction présente le plan d\'actionnariat salarié 2026, qui fait suite au plan mis en place en 2024. '
    'Lors de l\'édition 2024, 68 % des salariés avaient souscrit, ce qui avait été considéré comme un succès. '
    'Le plan 2026 est reconduit avec des modalités proches de celles de 2024, dans le prolongement du plan '
    'stratégique Vision 2030 du groupe. L\'évolution favorable du cours de l\'action rend ce nouveau plan '
    'particulièrement intéressant pour les salariés.'
)

add_heading2(doc, 'Modalités du plan')

add_normal_mixed(doc, [
    ('Bénéficiaires : ', True),
    ('Salariés en CDI et alternants présents dans les effectifs au 31 mars 2026. '
     'Les personnes en congé de mobilité encore présentes dans les effectifs à cette date peuvent également en bénéficier.', False),
])

add_normal_mixed(doc, [
    ('Décote : ', True),
    ('20 % sur le prix de référence de l\'action.', False),
])

add_normal_mixed(doc, [
    ('Abondement : ', True),
    ('100 % de l\'apport personnel du salarié, dans la limite de 600 euros bruts.', False),
])

add_normal_mixed(doc, [
    ('Prix de référence : ', True),
    ('Moyenne du cours de l\'action entre le 11 et le 30 mars 2026, à laquelle sera appliquée la décote de 20 %.', False),
])

add_normal_mixed(doc, [
    ('Facilités de paiement : ', True),
    ('Souscription directe ou déduction sur salaire en 12 mensualités (plafonnées à 10 % de la rémunération mensuelle), '
     'ou prélèvement sur compte bancaire.', False),
])

add_normal_mixed(doc, [
    ('Fiscalité : ', True),
    ('La décote est exonérée d\'impôt sur le revenu. L\'abondement est soumis à la CSG-CRDS (9,7 %). '
     'Les plus-values sont exonérées d\'impôt à la sortie du plan.', False),
])

add_normal_mixed(doc, [
    ('Blocage : ', True),
    ('5 ans, avec possibilité de déblocage anticipé dans les cas prévus par la loi '
     '(mariage, naissance, divorce, cessation du contrat de travail, etc.).', False),
])

add_normal_mixed(doc, [
    ('Dividendes : ', True),
    ('Réinvestis dans le Fonds Commun de Placement d\'Entreprise (FCPE).', False),
])

add_heading2(doc, 'Dates clés')

dates_table = doc.add_table(rows=5, cols=2)
dates_table.style = 'Table Grid'
data = [
    ('D\'ici le 17 février 2026', 'Signature de l\'avenant au PEE par les membres du CSE'),
    ('11 mars 2026', 'Fixation du prix de référence par le DG du groupe'),
    ('13 au 31 mars 2026', 'Période de souscription pour les salariés'),
    ('12 mai 2026', 'Confirmation des allocations et montants définitifs'),
    ('À noter', 'En cas de forte demande, des réductions d\'allocation peuvent s\'appliquer (comme en 2024)'),
]
for i, (k, v) in enumerate(data):
    dates_table.rows[i].cells[0].text = k
    dates_table.rows[i].cells[1].text = v

add_empty_para(doc)

add_heading2(doc, 'Discussions et remarques')

remarks = [
    'Les élus demandent des précisions sur l\'accès au portail de souscription pour les personnes en congé de mobilité '
    'n\'ayant plus accès à leur session professionnelle. La Direction s\'engage à remonter ce point aux équipes concernées '
    'et à envisager l\'utilisation d\'adresses email personnelles.',
    'Un débat a lieu sur la numérotation de l\'avenant au Plan d\'Épargne Entreprise (PEE). Après recherche, '
    'il s\'avère qu\'un premier avenant avait déjà été signé en 2023 (concernant le FCPE) et un deuxième en février 2024. '
    'L\'avenant actuel serait donc le n°3. La Direction vérifiera ce point.',
    'Les documents (présentation PowerPoint et avenant) seront transmis par email aux membres du CSE.',
]
for r in remarks:
    add_bullet(doc, r)

add_empty_para(doc)
add_normal_mixed(doc, [
    ('Vote sur l\'avis du CSE relatif au plan d\'actionnariat salarié 2026 :', True),
])

vote_table = doc.add_table(rows=4, cols=2)
vote_table.style = 'Table Grid'
vote_data = [('Pour', 'Unanimité'), ('Contre', '0'), ('Abstention', '0'), ('Résultat', 'AVIS FAVORABLE')]
for i, (k, v) in enumerate(vote_data):
    vote_table.rows[i].cells[0].text = k
    c = vote_table.rows[i].cells[1]
    c.text = ''
    run = c.paragraphs[0].add_run(v)
    if i == 3:
        run.bold = True
        vote_table.rows[i].cells[0].paragraphs[0].runs[0] if vote_table.rows[i].cells[0].paragraphs[0].runs else None
        # Bold the result row label too
        vote_table.rows[i].cells[0].text = ''
        run0 = vote_table.rows[i].cells[0].paragraphs[0].add_run('Résultat')
        run0.bold = True

add_empty_para(doc)

# ===== POINT 3 =====
add_heading1(doc, 'Point 3 – Consultation sur la rupture conventionnelle d\'une salariée protégée')

add_normal(doc,
    'La Direction informe le CSE de la procédure de rupture conventionnelle concernant une salariée protégée. '
    'Conformément à la législation, cette rupture nécessite :'
)

items = [
    'La consultation préalable du CSE (avis rendu ce jour)',
    'La demande d\'autorisation auprès de l\'Inspection du travail',
    'Un délai d\'instruction pouvant aller jusqu\'à 2 mois maximum',
]
for item in items:
    add_bullet(doc, item)

add_normal(doc,
    'La Direction précise que la même procédure sera suivie pour d\'autres cas à venir (notamment Clément). '
    'Les élus souhaitent que les démarches soient anticipées autant que possible afin de ne pas retarder '
    'les départs prévus. La Direction confirme que la consultation du CSE aujourd\'hui ne signifie pas '
    'une signature immédiate de la rupture, mais permet d\'engager le processus formel.'
)

add_empty_para(doc)
add_normal_mixed(doc, [
    ('Pièces transmises à l\'Inspection du travail : ', True),
    ('la convention signée, le CERFA, l\'avis du CSE.', False),
])

add_empty_para(doc)
add_normal_mixed(doc, [
    ('Vote sur la rupture conventionnelle de la salariée protégée :', True),
])

vote_table2 = doc.add_table(rows=4, cols=2)
vote_table2.style = 'Table Grid'
for i, (k, v) in enumerate(vote_data):
    vote_table2.rows[i].cells[0].text = k
    c = vote_table2.rows[i].cells[1]
    c.text = ''
    run = c.paragraphs[0].add_run(v)
    if i == 3:
        run.bold = True
        vote_table2.rows[i].cells[0].text = ''
        run0 = vote_table2.rows[i].cells[0].paragraphs[0].add_run('Résultat')
        run0.bold = True

add_empty_para(doc)

# ===== POINT 4 =====
add_heading1(doc, 'Point 4 – Suivi de l\'accompagnement des congés de mobilité (LHH)')

add_heading2(doc, 'Présentation du dispositif')
add_normal(doc,
    'Le représentant du cabinet LHH (Lee Hecht Harrison) présente le dispositif d\'accompagnement des salariés '
    'en congé de mobilité, faisant suite à la Rupture Conventionnelle Collective (RCC). '
    'Il s\'agit d\'une phase nouvelle, distincte de la phase de validation des projets (commissions de validation) '
    'qui a eu lieu précédemment.'
)

add_heading2(doc, 'Modalités d\'accompagnement')
modalities = [
    'Deux vagues de départ sont prévues : mars 2026 et juin 2026.',
    'Dès qu\'un salarié accepte le congé de mobilité et que les RH transmettent ses coordonnées, '
    'un consultant LHH dédié le contacte pour démarrer l\'accompagnement.',
    'Le même consultant suit le salarié pendant toute la durée du congé de mobilité.',
    'Durées du congé de mobilité : 6 mois pour les moins de 50 ans, 8 mois pour les plus de 50 ans, '
    'et jusqu\'à 10 mois pour les projets de création d\'entreprise.',
    'Possibilité de 2 mois supplémentaires sur décision de la commission de suivi.',
]
for m in modalities:
    add_bullet(doc, m)

add_heading2(doc, 'Types d\'accompagnement proposés')
types_acc = [
    'Recherche d\'emploi : aide au CV, préparation aux entretiens, accès aux bases de données d\'offres, coaching individuel.',
    'Formation qualifiante ou non qualifiante : suivi des inscriptions et démarches, aide au montage du dossier.',
    'Création d\'entreprise : ateliers dédiés (fortement recommandés).',
    'Ateliers collectifs (en présentiel à la Défense ou en distanciel) : storytelling, réseaux sociaux professionnels, pitch, etc.',
    'Entretiens individuels : en présentiel ou distanciel au choix du salarié.',
]
for t in types_acc:
    add_bullet(doc, t)

add_heading2(doc, 'Chiffres clés LHH')
add_normal(doc,
    'LHH accompagne environ 25 000 personnes par an et met en place 1 500 projets. '
    'En Île-de-France, le taux de reclassement est de 91 %, réparti comme suit :'
)
stats = ['34 % en emploi', '26 % en formation de reconversion avec projet d\'emploi', '23 % en création d\'entreprise']
for s in stats:
    add_bullet(doc, s)

add_heading2(doc, 'Discussions et points d\'attention soulevés par les élus')
discussions = [
    'Caractère non obligatoire des ateliers : les ateliers et entretiens proposés ne sont pas obligatoires. '
    'Cependant, les points réguliers avec le consultant sont importants pour maintenir un rythme et un suivi. '
    'Les élus insistent sur le fait que les salariés ne doivent pas être pénalisés s\'ils choisissent de ne pas '
    'suivre tous les ateliers, notamment s\'ils se sentent capables de gérer leur recherche de manière autonome.',
    'Restaurer la confiance : certains salariés ont exprimé une déception vis-à-vis de l\'accompagnement LHH '
    'lors de la phase précédente (période de décembre-janvier), jugée trop rapide et insuffisante. '
    'Les élus demandent que LHH fasse un effort pour rétablir le lien de confiance avec ces personnes, '
    'en reconnaissant que la période était contrainte et en démontrant la valeur ajoutée de la nouvelle phase d\'accompagnement.',
    'Commission de suivi : la commission de suivi (direction à voix décisionnaire + CSE à titre consultatif) décidera de l\'attribution '
    'des 2 mois supplémentaires en fin de congé de mobilité. Les élus demandent qu\'un contradictoire soit possible '
    'pour les salariés dont la prolongation serait refusée.',
    'Dates de départ : certaines formations commencent avant la date effective de début du congé de mobilité. '
    'Les élus alertent sur le décalage entre la date communiquée initialement (1er février) et la réalité, '
    'qui peut impacter les projets des salariés.',
    'Coordonnées personnelles : LHH vérifiera que les coordonnées transmises sont bien les coordonnées personnelles '
    'des salariés (et non professionnelles) pour assurer la continuité de la communication après le départ.',
]
for d in discussions:
    add_bullet(doc, d)

add_empty_para(doc)
add_normal_mixed(doc, [('Actions convenues :', True)])
actions = [
    'La Direction enverra un message aux salariés concernés pour présenter la nouvelle phase d\'accompagnement, '
    'en mettant en avant les chiffres de reclassement et la valeur du dispositif.',
    'Les consultants LHH recontacteront chaque salarié dès la signature de leur congé de mobilité.',
    'Un point de suivi sera refait en CSE dans environ 2 mois (avril 2026).',
    'La Direction vérifiera les coordonnées personnelles des salariés transmises à LHH.',
]
for a in actions:
    add_bullet(doc, a)

add_empty_para(doc)

# ===== POINT 5 =====
add_heading1(doc, 'Point 5 – Point sur le projet Career Pass')

add_normal(doc,
    'Loïc, nouveau Responsable RH en charge de la R&D France (depuis le 1er janvier 2026), '
    'rejoint la réunion en visioconférence pour présenter l\'état d\'avancement '
    'du projet Career Pass (outil de gestion des parcours professionnels et carrières).'
)

add_heading2(doc, 'État des lieux')
etat = [
    'Un travail considérable a été réalisé par les équipes Netatmo pour construire un outil de modélisation des fonctions '
    'et des parcours de carrière. Loïc salue la qualité et la granularité de ce travail.',
    'L\'outil permet de clarifier les évolutions possibles, les attendus managériaux, et peut contribuer à l\'attractivité '
    'et à la fidélisation des collaborateurs.',
    'Cependant, cet outil doit être articulé avec le projet de cartographie des emplois en cours au niveau du groupe Legrand.',
]
for e in etat:
    add_bullet(doc, e)

add_heading2(doc, 'Points de vigilance')
vigilance = [
    'Homogénéisation avec le groupe : le travail de Netatmo est à un niveau de granularité (fonctions) plus fin que '
    'celui du groupe (emplois au sens de la convention collective de la métallurgie). Un travail de transposition est nécessaire.',
    'Nouvelles entités : l\'arrivée de Cogérec et les synergies avec d\'autres entités du groupe (Barco, etc.) '
    'rendent cette homogénéisation d\'autant plus importante.',
    'Intégration dans Harmonie (outil SIRH du groupe) : il faut vérifier la compatibilité technique et s\'assurer '
    'que la granularité spécifique aux métiers Netatmo (software, cyber, qualité, etc.) ne soit pas perdue dans '
    'l\'harmonisation, comme cela a pu être le cas pour certaines fonctionnalités des entretiens individuels.',
]
for v in vigilance:
    add_bullet(doc, v)

add_heading2(doc, 'Calendrier envisagé')
calendar_items = [
    'Des tests seront réalisés chez Netatmo (sur 1 ou 2 métiers : software ou R&D) pour vérifier la cohérence du Career Pass '
    'avec le référentiel groupe.',
    'Échéance visée : courant 2026, probablement après l\'été (plutôt second semestre).',
    'Loïc s\'engage à respecter ses engagements de calendrier et préfère être prudent plutôt que de surpromettre.',
]
for c in calendar_items:
    add_bullet(doc, c)

add_heading2(doc, 'Remarques des élus')
add_normal(doc,
    'Les élus expriment une vive préoccupation concernant les délais : le Career Pass est attendu depuis plusieurs années '
    'par les collaborateurs (certains évoquent 3 à 4 ans). Les entretiens individuels remontent systématiquement cette attente. '
    'Les élus alertent sur le risque de déception et de désengagement si le projet est encore repoussé, '
    'et sur le sentiment que l\'inertie du groupe Legrand ralentit les initiatives locales de Netatmo. '
    'La Direction et Loïc assurent que le sujet est bien pris en compte, que le projet sera traité en 2026, '
    'et qu\'il ne s\'agit pas de procrastination mais d\'un travail nécessaire d\'alignement.'
)

add_empty_para(doc)

# ===== POINT 6 =====
add_heading1(doc, 'Point 6 – Situation de l\'équipe Office Management – Alerte RPS')

add_normal(doc,
    'Un élu rapporte qu\'Anne Savard, de l\'équipe Office Management, a contacté le CSE pour signaler '
    'une situation de mal-être au travail. Les éléments rapportés incluent :'
)
rps_items = [
    'Un sentiment d\'imposition de décisions sans concertation',
    'Une dévalorisation de son poste',
    'Un manque de reconnaissance',
    'Des formes d\'isolement',
    'Une situation pouvant relever des Risques Psychosociaux (RPS)',
]
for r in rps_items:
    add_bullet(doc, r)

add_normal(doc, 'Anne Savard a été absente plus de 5 mois en 2025 en lien avec cette situation.')

add_empty_para(doc)
add_normal_mixed(doc, [('Actions en cours :', True)])
actions_rps = [
    'La Direction a proposé un avenant au contrat de travail avec ajustements horaires (passage à 80 %), '
    'en coordination avec la médecine du travail et le conseil juridique, mais se heurte à une résistance de la salariée.',
    'Une enquête RPS a été déclenchée par la Direction, incluant une inspection de l\'environnement de travail '
    'de l\'équipe et des entretiens individuels avec les membres de l\'équipe (sur base volontaire).',
    'L\'enquête est menée de manière paritaire. Les résultats seront partagés avec le CSE.',
    'La Direction indique avoir des éléments solides à partager mais souhaite attendre les conclusions de l\'enquête '
    'avant d\'engager des actions supplémentaires.',
]
for a in actions_rps:
    add_bullet(doc, a)

add_empty_para(doc)
add_normal(doc,
    'Les élus demandent à être associés à l\'enquête et rappellent que l\'inspection est une prérogative du CSE. '
    'Un vote du CSE pourrait être nécessaire pour formaliser le déclenchement de cette inspection. '
    'La Direction confirme sa volonté de transparence et de coordination avec le CSE sur ce dossier.'
)

add_empty_para(doc)

# ===== POINT 7 =====
add_heading1(doc, 'Point 7 – Questions diverses')
add_normal(doc, 'Aucune question diverse supplémentaire n\'a été soulevée.')

add_empty_para(doc)

# ===== CLOTURE =====
add_heading1(doc, 'Clôture de la séance')
add_normal(doc,
    'L\'ordre du jour étant épuisé, la séance est levée. '
    'La prochaine réunion du CSE est prévue courant mars 2026.'
)

add_empty_para(doc, 2)

# ===== SIGNATURES =====
p = add_normal(doc)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Signatures')
run.underline = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 0, 0)

add_empty_para(doc)

sig_table = doc.add_table(rows=3, cols=2)
sig_table.style = 'Table Grid'
sig_table.rows[0].cells[0].text = 'Alexandre MENU, Président'
sig_table.rows[0].cells[1].text = ''
sig_table.rows[1].cells[0].text = 'Sarah ETCHEVERRY, Direction'
sig_table.rows[1].cells[1].text = ''
sig_table.rows[2].cells[0].text = 'Erwan PINAULT, Secrétaire'
sig_table.rows[2].cells[1].text = ''

add_empty_para(doc, 2)

# Final heading
add_heading1(doc, 'Réunion CSE du 10/02/26')
add_heading2(doc, '')

# Save
output_path = '/Users/epinault/tmp/CSE/PV_CSE_10_fevrier_2026.docx'
doc.save(output_path)
print(f'PV généré avec succès au format template : {output_path}')
