from docx import Document
from docx.shared import Pt, Inches, Emu
from docx.oxml.ns import qn
import json

doc = Document("/Users/epinault/tmp/CSE/template CSE.docx")

print("=" * 80)
print("TEMPLATE CSE - INSPECTION")
print("=" * 80)

# Sections
print("\n--- SECTIONS ---")
for i, section in enumerate(doc.sections):
    print(f"Section {i}:")
    print(f"  Page width: {section.page_width}, height: {section.page_height}")
    print(f"  Margins: top={section.top_margin}, bottom={section.bottom_margin}, left={section.left_margin}, right={section.right_margin}")
    print(f"  Header distance: {section.header_distance}, Footer distance: {section.footer_distance}")
    # Check header/footer
    header = section.header
    if header and not header.is_linked_to_previous:
        print(f"  Header has content: {len(header.paragraphs)} paragraphs")
        for hp in header.paragraphs:
            print(f"    Header para: '{hp.text}' align={hp.alignment}")
            for run in hp.runs:
                print(f"      Run: '{run.text}' bold={run.bold} italic={run.italic} size={run.font.size} color={run.font.color.rgb if run.font.color and run.font.color.rgb else 'none'}")
    footer = section.footer
    if footer and not footer.is_linked_to_previous:
        print(f"  Footer has content: {len(footer.paragraphs)} paragraphs")
        for fp in footer.paragraphs:
            print(f"    Footer para: '{fp.text}' align={fp.alignment}")

# Styles used
print("\n--- STYLES USED ---")
styles_seen = set()
for para in doc.paragraphs:
    styles_seen.add(para.style.name)
print(f"Paragraph styles: {sorted(styles_seen)}")

# Content
print("\n--- PARAGRAPHS ---")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    indent_info = ""
    if para.paragraph_format.left_indent:
        indent_info = f" indent={para.paragraph_format.left_indent}"
    space_info = ""
    if para.paragraph_format.space_before:
        space_info += f" space_before={para.paragraph_format.space_before}"
    if para.paragraph_format.space_after:
        space_info += f" space_after={para.paragraph_format.space_after}"

    print(f"[{i}] style='{para.style.name}' align={para.alignment}{indent_info}{space_info}")
    if text:
        print(f"     text: '{text[:120]}{'...' if len(text) > 120 else ''}'")
    for j, run in enumerate(para.runs):
        rtext = run.text.strip()
        if rtext or run.bold or run.italic or run.underline:
            font_info = f"bold={run.bold} italic={run.italic} underline={run.underline}"
            if run.font.size:
                font_info += f" size={run.font.size}"
            if run.font.name:
                font_info += f" font={run.font.name}"
            if run.font.color and run.font.color.rgb:
                font_info += f" color={run.font.color.rgb}"
            print(f"     run[{j}]: '{rtext[:80]}' {font_info}")

# Tables
print("\n--- TABLES ---")
for i, table in enumerate(doc.tables):
    print(f"Table {i}: {len(table.rows)} rows x {len(table.columns)} cols")
    print(f"  Style: {table.style.name if table.style else 'None'}")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            if cell_text:
                print(f"  [{ri},{ci}]: '{cell_text[:80]}'")
            # Check cell formatting
            for cp in cell.paragraphs:
                for cr in cp.runs:
                    if cr.font.size or cr.bold or cr.font.color.rgb if cr.font.color else None:
                        info = []
                        if cr.bold: info.append("bold")
                        if cr.font.size: info.append(f"size={cr.font.size}")
                        if cr.font.color and cr.font.color.rgb: info.append(f"color={cr.font.color.rgb}")
                        if cr.font.name: info.append(f"font={cr.font.name}")
                        if info:
                            print(f"         format: {', '.join(info)}")

# Check for images in header
print("\n--- HEADER IMAGES ---")
for i, section in enumerate(doc.sections):
    header = section.header
    if header:
        for rel_id, rel in header.part.rels.items():
            if "image" in rel.reltype:
                print(f"  Header image found: {rel.target_ref}")

# Check body images
print("\n--- BODY IMAGES ---")
for i, para in enumerate(doc.paragraphs):
    for run in para.runs:
        drawing_elements = run._element.findall(qn('w:drawing'))
        inline_elements = run._element.findall('.//' + qn('wp:inline'))
        if drawing_elements or inline_elements:
            print(f"  Para [{i}]: Image found")

# Check numbering
print("\n--- NUMBERING (list styles) ---")
for i, para in enumerate(doc.paragraphs):
    numPr = para._element.find(qn('w:pPr'))
    if numPr is not None:
        num = numPr.find(qn('w:numPr'))
        if num is not None:
            ilvl = num.find(qn('w:ilvl'))
            numId = num.find(qn('w:numId'))
            print(f"  Para [{i}]: numId={numId.get(qn('w:val')) if numId is not None else 'N/A'} ilvl={ilvl.get(qn('w:val')) if ilvl is not None else 'N/A'} text='{para.text[:60]}'")
