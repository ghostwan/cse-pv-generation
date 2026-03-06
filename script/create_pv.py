from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- EN-TETE ---
title = doc.add_heading('PROCES-VERBAL DE REUNION', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_heading('Comite Social et Economique (CSE)', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph('')

# --- INFORMATIONS GENERALES ---
info = doc.add_table(rows=4, cols=2)
info.style = 'Light Grid Accent 1'
info.alignment = WD_TABLE_ALIGNMENT.CENTER

cells = info.rows[0].cells
cells[0].text = 'Date de la reunion'
cells[1].text = 'Lundi 10 fevrier 2026'

cells = info.rows[1].cells
cells[0].text = 'Lieu'
cells[1].text = 'Locaux de l\'entreprise (+ visioconference)'

cells = info.rows[2].cells
cells[0].text = 'Duree'
cells[1].text = 'Environ 2h15'

cells = info.rows[3].cells
cells[0].text = 'Secretaire de seance'
cells[1].text = 'Erwan'

for row in info.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(10)

doc.add_paragraph('')

# --- PARTICIPANTS ---
doc.add_heading('Participants', level=2)

p = doc.add_paragraph()
p.add_run('Membres presents du CSE : ').bold = True
p.add_run('Erwan (secretaire de seance), Emile, Alice, Celine, Mickael, Brilleau, Jada, Nadege, Alexandre')

p = doc.add_paragraph()
p.add_run('Representants de la Direction : ').bold = True
p.add_run('Marc (President du CSE)')

p = doc.add_paragraph()
p.add_run('Invites : ').bold = True
p.add_run('Loic (Responsable RH R&D France - en visioconference), representant LHH (cabinet d\'accompagnement)')

p = doc.add_paragraph()
p.add_run('Absents excuses : ').bold = True
p.add_run('Christophe (membre du CSE, en arret), Antoine (a participe partiellement)')

doc.add_paragraph('')

# --- ORDRE DU JOUR ---
doc.add_heading('Ordre du jour', level=2)

items = [
    'Approbation du PV de la reunion precedente',
    'Consultation sur le plan d\'actionnariat salarie 2026',
    'Consultation sur la rupture conventionnelle d\'une salariee protegee',
    'Suivi de l\'accompagnement des conges de mobilite (LHH)',
    'Point sur le projet Career Pass (carrieres et parcours professionnels)',
    'Situation de l\'equipe Office Management - Alerte RPS',
    'Questions diverses',
]
for i, item in enumerate(items, 1):
    doc.add_paragraph(f'{i}. {item}', style='List Number')

doc.add_paragraph('')
doc.add_paragraph('').add_run('_' * 70)

# ============================
# POINT 1
# ============================
doc.add_heading('1. Approbation du PV de la reunion precedente', level=2)
p = doc.add_paragraph()
p.add_run('Le proces-verbal de la reunion precedente est soumis a l\'approbation des membres du CSE.')
doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Decision : ').bold = True
p.add_run('Le PV est approuve.')

doc.add_paragraph('')
doc.add_paragraph('').add_run('_' * 70)

# ============================
# POINT 2
# ============================
doc.add_heading('2. Consultation sur le plan d\'actionnariat salarie 2026', level=2)

doc.add_heading('Contexte', level=3)
p = doc.add_paragraph()
p.add_run(
    'La Direction presente le plan d\'actionnariat salarie 2026, qui fait suite au plan mis en place en 2024. '
    'Lors de l\'edition 2024, 68 % des salaries avaient souscrit, ce qui avait ete considere comme un succes. '
    'Le plan 2026 est reconduit avec des modalites proches de celles de 2024, dans le prolongement du plan '
    'strategique Vision 2030 du groupe. L\'evolution favorable du cours de l\'action rend ce nouveau plan '
    'particulierement interessant pour les salaries.'
)

doc.add_heading('Modalites du plan', level=3)
p = doc.add_paragraph()
p.add_run('Beneficiaires : ').bold = True
p.add_run(
    'Salaries en CDI et alternants presents dans les effectifs au 31 mars 2026. '
    'Les personnes en conge de mobilite encore presentes dans les effectifs a cette date peuvent egalement en beneficier.'
)

p = doc.add_paragraph()
p.add_run('Decote : ').bold = True
p.add_run('20 % sur le prix de reference de l\'action.')

p = doc.add_paragraph()
p.add_run('Abondement : ').bold = True
p.add_run('100 % de l\'apport personnel du salarie, dans la limite de 600 euros bruts.')

p = doc.add_paragraph()
p.add_run('Prix de reference : ').bold = True
p.add_run('Moyenne du cours de l\'action entre le 11 et le 30 mars 2026, a laquelle sera appliquee la decote de 20 %.')

p = doc.add_paragraph()
p.add_run('Facilites de paiement : ').bold = True
p.add_run(
    'Souscription directe ou deduction sur salaire en 12 mensualites (plafonnees a 10 % de la remuneration mensuelle), '
    'ou prelevement sur compte bancaire.'
)

p = doc.add_paragraph()
p.add_run('Fiscalite : ').bold = True
p.add_run(
    'La decote est exoneree d\'impot sur le revenu. L\'abondement est soumis a la CSG-CRDS (9,7 %). '
    'Les plus-values sont exonerees d\'impot a la sortie du plan.'
)

p = doc.add_paragraph()
p.add_run('Blocage : ').bold = True
p.add_run(
    '5 ans, avec possibilite de deblocage anticipe dans les cas prevus par la loi '
    '(mariage, naissance, divorce, cessation du contrat de travail, etc.).'
)

p = doc.add_paragraph()
p.add_run('Dividendes : ').bold = True
p.add_run('Reinvestis dans le Fonds Commun de Placement d\'Entreprise (FCPE).')

doc.add_heading('Dates cles', level=3)

dates_table = doc.add_table(rows=5, cols=2)
dates_table.style = 'Light Grid Accent 1'
dates_table.rows[0].cells[0].text = 'D\'ici le 17 fevrier 2026'
dates_table.rows[0].cells[1].text = 'Signature de l\'avenant au PEE par les membres du CSE'
dates_table.rows[1].cells[0].text = '11 mars 2026'
dates_table.rows[1].cells[1].text = 'Fixation du prix de reference par le DG du groupe'
dates_table.rows[2].cells[0].text = '13 au 31 mars 2026'
dates_table.rows[2].cells[1].text = 'Periode de souscription pour les salaries'
dates_table.rows[3].cells[0].text = '12 mai 2026'
dates_table.rows[3].cells[1].text = 'Confirmation des allocations et montants definitifs'
dates_table.rows[4].cells[0].text = 'A noter'
dates_table.rows[4].cells[1].text = 'En cas de forte demande, des reductions d\'allocation peuvent s\'appliquer (comme en 2024)'

doc.add_paragraph('')

doc.add_heading('Discussions et remarques', level=3)
points = [
    'Les elus demandent des precisions sur l\'acces au portail de souscription pour les personnes en conge de mobilite '
    'n\'ayant plus acces a leur session professionnelle. La Direction s\'engage a remonter ce point aux equipes concernees '
    'et a envisager l\'utilisation d\'adresses email personnelles.',
    'Un debat a lieu sur la numerotation de l\'avenant au Plan d\'Epargne Entreprise (PEE). Apres recherche, '
    'il s\'avere qu\'un premier avenant avait deja ete signe en 2023 (concernant le FCPE) et un deuxieme en fevrier 2024. '
    'L\'avenant actuel serait donc le n\u00b03. La Direction verifiera ce point.',
    'Les documents (presentation PowerPoint et avenant) seront transmis par email aux membres du CSE.',
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Vote sur l\'avis du CSE relatif au plan d\'actionnariat salarie 2026 :').bold = True
doc.add_paragraph('')

vote_table = doc.add_table(rows=4, cols=2)
vote_table.style = 'Light Grid Accent 1'
vote_table.rows[0].cells[0].text = 'Pour'
vote_table.rows[0].cells[1].text = 'Unanimite'
vote_table.rows[1].cells[0].text = 'Contre'
vote_table.rows[1].cells[1].text = '0'
vote_table.rows[2].cells[0].text = 'Abstention'
vote_table.rows[2].cells[1].text = '0'
vote_table.rows[3].cells[0].text = 'Resultat'
vote_table.rows[3].cells[1].text = 'AVIS FAVORABLE'
for cell in vote_table.rows[3].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph('')
doc.add_paragraph('').add_run('_' * 70)

# ============================
# POINT 3
# ============================
doc.add_heading('3. Consultation sur la rupture conventionnelle d\'une salariee protegee', level=2)

p = doc.add_paragraph()
p.add_run(
    'La Direction informe le CSE de la procedure de rupture conventionnelle concernant une salariee protegee. '
    'Conformement a la legislation, cette rupture necessite :'
)

items = [
    'La consultation prealable du CSE (avis rendu ce jour)',
    'La demande d\'autorisation aupres de l\'Inspection du travail',
    'Un delai d\'instruction pouvant aller jusqu\'a 2 mois maximum',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

p = doc.add_paragraph()
p.add_run(
    'La Direction precise que la meme procedure sera suivie pour d\'autres cas a venir (notamment Clement). '
    'Les elus souhaitent que les demarches soient anticipees autant que possible afin de ne pas retarder '
    'les departs prevus. La Direction confirme que la consultation du CSE aujourd\'hui ne signifie pas '
    'une signature immediate de la rupture, mais permet d\'engager le processus formel.'
)

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Les pieces transmises a l\'Inspection du travail comprendront : ').bold = True
p.add_run('la convention signee, le CERFA, l\'avis du CSE.')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Vote sur la rupture conventionnelle de la salariee protegee :').bold = True

vote_table2 = doc.add_table(rows=4, cols=2)
vote_table2.style = 'Light Grid Accent 1'
vote_table2.rows[0].cells[0].text = 'Pour'
vote_table2.rows[0].cells[1].text = 'Unanimite'
vote_table2.rows[1].cells[0].text = 'Contre'
vote_table2.rows[1].cells[1].text = '0'
vote_table2.rows[2].cells[0].text = 'Abstention'
vote_table2.rows[2].cells[1].text = '0'
vote_table2.rows[3].cells[0].text = 'Resultat'
vote_table2.rows[3].cells[1].text = 'AVIS FAVORABLE'
for cell in vote_table2.rows[3].cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

doc.add_paragraph('')
doc.add_paragraph('').add_run('_' * 70)

# ============================
# POINT 4
# ============================
doc.add_heading('4. Suivi de l\'accompagnement des conges de mobilite (cabinet LHH)', level=2)

doc.add_heading('Presentation du dispositif', level=3)
p = doc.add_paragraph()
p.add_run(
    'Le representant du cabinet LHH (Lee Hecht Harrison) presente le dispositif d\'accompagnement des salaries '
    'en conge de mobilite, faisant suite a la Rupture Conventionnelle Collective (RCC). '
    'Il s\'agit d\'une phase nouvelle, distincte de la phase de validation des projets (commissions de validation) '
    'qui a eu lieu precedemment.'
)

doc.add_heading('Modalites d\'accompagnement', level=3)
points = [
    'Deux vagues de depart sont prevues : mars 2026 et juin 2026.',
    'Des qu\'un salarie accepte le conge de mobilite et que les RH transmettent ses coordonnees, '
    'un consultant LHH dedie le contacte pour demarrer l\'accompagnement.',
    'Le meme consultant suit le salarie pendant toute la duree du conge de mobilite.',
    'Durees du conge de mobilite : 6 mois pour les moins de 50 ans, 8 mois pour les plus de 50 ans, '
    'et jusqu\'a 10 mois pour les projets de creation d\'entreprise.',
    'Possibilite de 2 mois supplementaires sur decision de la commission de suivi.',
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Types d\'accompagnement proposes', level=3)
points = [
    'Recherche d\'emploi : aide au CV, preparation aux entretiens, acces aux bases de donnees d\'offres, coaching individuel.',
    'Formation qualifiante ou non qualifiante : suivi des inscriptions et demarches, aide au montage du dossier.',
    'Creation d\'entreprise : ateliers dedies (fortement recommandes).',
    'Ateliers collectifs (en presentiel, a la Defense ou en distanciel) : storytelling, reseaux sociaux professionnels, pitch, etc.',
    'Entretiens individuels : en presentiel ou distanciel au choix du salarie.',
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Chiffres cles LHH', level=3)
p = doc.add_paragraph()
p.add_run(
    'LHH accompagne environ 25 000 personnes par an et met en place 1 500 projets. '
    'En Ile-de-France, le taux de reclassement est de 91 %, reparti comme suit :'
)
points = [
    '34 % en emploi',
    '26 % en formation de reconversion avec projet d\'emploi',
    '23 % en creation d\'entreprise',
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Discussions et points d\'attention souleves par les elus', level=3)
points = [
    'Caractere non obligatoire des ateliers : les ateliers et entretiens proposes ne sont pas obligatoires. '
    'Cependant, les points reguliers avec le consultant sont importants pour maintenir un rythme et un suivi. '
    'Les elus insistent sur le fait que les salaries ne doivent pas etre penalises s\'ils choisissent de ne pas '
    'suivre tous les ateliers, notamment s\'ils se sentent capables de gerer leur recherche de maniere autonome.',
    'Restaurer la confiance : certains salaries ont exprime une deception vis-a-vis de l\'accompagnement LHH '
    'lors de la phase precedente (periode de decembre-janvier), jugee trop rapide et insuffisante. '
    'Les elus demandent que LHH fasse un effort pour retablir le lien de confiance avec ces personnes, '
    'en reconnaissant que la periode etait contrainte et en demontrant la valeur ajoutee de la nouvelle phase d\'accompagnement.',
    'Commission de suivi : la commission de suivi (direction + CSE a titre consultatif) decidera de l\'attribution '
    'des 2 mois supplementaires en fin de conge de mobilite. Les elus demandent qu\'un contradictoire soit possible '
    'pour les salaries dont la prolongation serait refusee.',
    'Dates de depart : certaines formations commencent avant la date effective de debut du conge de mobilite. '
    'Les elus alertent sur le decalage entre la date communiquee initialement (1er fevrier) et la realite, '
    'qui peut impacter les projets des salaries.',
    'Coordonnees personnelles : LHH verifiera que les coordonnees transmises sont bien les coordonnees personnelles '
    'des salaries (et non professionnelles) pour assurer la continuite de la communication apres le depart.',
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Actions convenues : ').bold = True
actions = [
    'La Direction enverra un message aux salaries concernes pour presenter la nouvelle phase d\'accompagnement, '
    'en mettant en avant les chiffres de reclassement et la valeur du dispositif.',
    'Les consultants LHH recontacteront chaque salarie des la signature de leur conge de mobilite.',
    'Un point de suivi sera refait en CSE dans environ 2 mois (avril 2026).',
    'La Direction verifiera les coordonnees personnelles des salaries transmises a LHH.',
]
for action in actions:
    doc.add_paragraph(action, style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('').add_run('_' * 70)

# ============================
# POINT 5
# ============================
doc.add_heading('5. Point sur le projet Career Pass', level=2)

p = doc.add_paragraph()
p.add_run(
    'Loic, nouveau Responsable RH en charge de la R&D France (depuis le 1er janvier 2026, succedant a un perimetre '
    'anterieurement centre sur la SBUI), rejoint la reunion en visioconference pour presenter l\'etat d\'avancement '
    'du projet Career Pass (outil de gestion des parcours professionnels et carrieres).'
)

doc.add_heading('Etat des lieux', level=3)
points = [
    'Un travail considerable a ete realise par les equipes Netatmo pour construire un outil de modelisation des fonctions '
    'et des parcours de carriere. Loic salue la qualite et la granularite de ce travail.',
    'L\'outil permet de clarifier les evolutions possibles, les attendus manageants, et peut contribuer a l\'attractivite '
    'et a la fidelisation des collaborateurs.',
    'Cependant, cet outil doit etre articule avec le projet de cartographie des emplois en cours au niveau du groupe Legrand.',
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Points de vigilance', level=3)
points = [
    'Homogeneisation avec le groupe : le travail de Netatmo est a un niveau de granularite (fonctions) plus fin que '
    'celui du groupe (emplois au sens de la convention collective de la metallurgie). Un travail de transposition est necessaire.',
    'Nouvelles entites : l\'arrivee de Cogerec et les synergies avec d\'autres entites du groupe (Barco, etc.) '
    'rendent cette homogeneisation d\'autant plus importante.',
    'Integration dans Harmonie (outil SIRH du groupe) : il faut verifier la compatibilite technique et s\'assurer '
    'que la granularite specifique aux metiers Netatmo (software, cyber, qualite, etc.) ne soit pas perdue dans '
    'l\'harmonisation, comme cela a pu etre le cas pour certaines fonctionnalites des entretiens individuels.',
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Calendrier envisage', level=3)
points = [
    'Des tests seront realises chez Netatmo (sur 1 ou 2 metiers : software ou R&D) pour verifier la coherence du Career Pass '
    'avec le referentiel groupe.',
    'Echeance visee : courant 2026, probablement apres l\'ete (plutot second semestre).',
    'Loic s\'engage a respecter ses engagements de calendrier et prefere etre prudent plutot que de surpromettre.',
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('Remarques des elus', level=3)
p = doc.add_paragraph()
p.add_run(
    'Les elus expriment une vive preoccupation concernant les delais : le Career Pass est attendu depuis plusieurs annees '
    'par les collaborateurs (certains evoquent 3 a 4 ans). Les entretiens individuels remontent systematiquement cette attente. '
    'Les elus alertent sur le risque de deception et de desengagement si le projet est encore repousse, '
    'et sur le sentiment que l\'inertie du groupe Legrand ralentit les initiatives locales de Netatmo. '
    'La Direction et Loic assurent que le sujet est bien pris en compte, que le projet sera traite en 2026, '
    'et qu\'il ne s\'agit pas de procrastination mais d\'un travail necessaire d\'alignement.'
)

doc.add_paragraph('')
doc.add_paragraph('').add_run('_' * 70)

# ============================
# POINT 6
# ============================
doc.add_heading('6. Situation de l\'equipe Office Management - Alerte RPS', level=2)

p = doc.add_paragraph()
p.add_run(
    'Un elu rapporte qu\'Anne Savard, de l\'equipe Office Management, a contacte le CSE pour signaler '
    'une situation de mal-etre au travail. Les elements rapportes incluent :'
)
points = [
    'Un sentiment d\'imposition de decisions sans concertation',
    'Une devalorisation de son poste',
    'Un manque de reconnaissance',
    'Des formes d\'isolement',
    'Une situation pouvant relever des Risques Psychosociaux (RPS)',
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

p = doc.add_paragraph()
p.add_run(
    'Anne Savard a ete absente plus de 5 mois en 2025 en lien avec cette situation.'
)

doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('Actions en cours : ').bold = True
points = [
    'La Direction a propose un avenant au contrat de travail (ajustements horaires) en coordination avec '
    'la medecine du travail et le conseil juridique, mais se heurte a une resistance de la salariee.',
    'Une enquete RPS a ete declenchee par la Direction, incluant une inspection de l\'environnement de travail '
    'de l\'equipe et des entretiens individuels avec les membres de l\'equipe (sur base volontaire).',
    'L\'enquete est menee de maniere paritaire. Les resultats seront partages avec le CSE.',
    'La Direction indique avoir des elements a partager mais souhaite attendre les conclusions de l\'enquete '
    'avant d\'engager des actions supplementaires.',
]
for point in points:
    doc.add_paragraph(point, style='List Bullet')

p = doc.add_paragraph()
p.add_run(
    'Les elus demandent a etre associes a l\'enquete et rappellent que l\'inspection est une prerogative du CSE. '
    'Un vote du CSE pourrait etre necessaire pour formaliser le declenchement de cette inspection. '
    'La Direction confirme sa volonte de transparence et de coordination avec le CSE sur ce dossier.'
)

doc.add_paragraph('')
doc.add_paragraph('').add_run('_' * 70)

# ============================
# POINT 7
# ============================
doc.add_heading('7. Questions diverses', level=2)
p = doc.add_paragraph()
p.add_run('Aucune question diverse supplementaire n\'a ete soulevee.')

doc.add_paragraph('')
doc.add_paragraph('').add_run('_' * 70)

# ============================
# CLOTURE
# ============================
doc.add_heading('Cloture de la seance', level=2)
p = doc.add_paragraph()
p.add_run(
    'L\'ordre du jour etant epuise, la seance est levee. '
    'La prochaine reunion du CSE est prevue courant mars 2026.'
)

doc.add_paragraph('')
doc.add_paragraph('')

# Signatures
sig_table = doc.add_table(rows=2, cols=2)
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sig_table.rows[0].cells[0].text = 'Le President du CSE'
sig_table.rows[0].cells[1].text = 'Le Secretaire du CSE'
sig_table.rows[1].cells[0].text = '\n\n\nMarc'
sig_table.rows[1].cells[1].text = '\n\n\nErwan'

for row in sig_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- FOOTER ---
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Document confidentiel - Usage interne uniquement')
run.font.size = Pt(8)
run.font.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)

# Sauvegarder
output_path = '/Users/epinault/tmp/CSE/PV_CSE_10_fevrier_2026.docx'
doc.save(output_path)
print(f'PV genere avec succes : {output_path}')
